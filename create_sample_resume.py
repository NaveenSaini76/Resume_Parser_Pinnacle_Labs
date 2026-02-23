"""
create_sample_resume.py
=======================
Run this script ONCE to generate a sample resume DOCX file you can
use to test the Resume Parser AI application.

Usage:
    python create_sample_resume.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def add_heading(doc, text, level=1, color=(30, 64, 175)):
    """Add a styled heading paragraph."""
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.runs[0] if para.runs else para.add_run(text)
    run.font.color.rgb = RGBColor(*color)
    return para


def add_section_divider(doc):
    """Add a horizontal divider paragraph."""
    p = doc.add_paragraph('─' * 65)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(148, 163, 184)


def create_sample_resume():
    """Build and save a sample resume DOCX file."""
    doc = Document()

    # ------- HEADER -------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Alex Johnson")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 64, 175)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(
        "📧 alex.johnson@email.com  |  📞 +1 (555) 867-5309  |  "
        "🔗 https://www.linkedin.com/in/alex-johnson-dev  |  "
        "📍 San Francisco, CA"
    ).font.size = Pt(10)

    add_section_divider(doc)

    # ------- SUMMARY -------
    add_heading(doc, "Summary", level=2)
    doc.add_paragraph(
        "Results-driven Full Stack Software Engineer with 4+ years of experience "
        "designing and building scalable web applications. Proficient in Python, "
        "React, and cloud infrastructure (AWS). Passionate about clean code, "
        "CI/CD best practices, and machine learning integration."
    )

    add_section_divider(doc)

    # ------- SKILLS -------
    add_heading(doc, "Skills", level=2)
    doc.add_paragraph(
        "Programming Languages: Python, JavaScript, TypeScript, Java, SQL, Bash\n"
        "Frontend: React, Angular, HTML, CSS, Bootstrap, Tailwind\n"
        "Backend: Flask, Django, FastAPI, Node.js, Express, REST API, GraphQL\n"
        "Databases: PostgreSQL, MySQL, MongoDB, Redis, SQLite\n"
        "Cloud & DevOps: AWS, Docker, Kubernetes, Jenkins, Git, GitHub, CI/CD, Terraform\n"
        "Data Science: Pandas, NumPy, Scikit-Learn, TensorFlow, Machine Learning, NLP\n"
        "Tools: Jira, VS Code, Figma, Postman, Jupyter"
    )

    add_section_divider(doc)

    # ------- EXPERIENCE -------
    add_heading(doc, "Experience", level=2)

    # Job 1
    p = doc.add_paragraph()
    p.add_run("Senior Software Engineer — TechNova Inc., San Francisco, CA").bold = True
    doc.add_paragraph("June 2022 – Present")
    doc.add_paragraph(
        "• Led a team of 5 engineers to build a microservices platform using Python (FastAPI) "
        "and Docker deployed on AWS ECS.\n"
        "• Reduced API latency by 40% through Redis caching and query optimisation.\n"
        "• Implemented CI/CD pipelines with Jenkins and GitHub Actions, cutting release "
        "times from 2 days to 2 hours.\n"
        "• Mentored 3 junior developers and conducted weekly code-review sessions."
    )

    # Job 2
    p = doc.add_paragraph()
    p.add_run("Software Engineer — DataBridge Solutions, Austin, TX").bold = True
    doc.add_paragraph("July 2020 – May 2022")
    doc.add_paragraph(
        "• Developed full-stack web apps with React (frontend) and Django REST Framework (backend).\n"
        "• Migrated legacy monolith to microservices, improving system uptime to 99.95%.\n"
        "• Built real-time dashboards using WebSockets, Celery, and PostgreSQL.\n"
        "• Wrote comprehensive unit and integration tests, achieving 90%+ coverage."
    )

    add_section_divider(doc)

    # ------- EDUCATION -------
    add_heading(doc, "Education", level=2)
    p = doc.add_paragraph()
    p.add_run("B.Tech in Computer Science — Stanford University, CA").bold = True
    doc.add_paragraph("Graduated: May 2020  |  GPA: 3.8 / 4.0")
    doc.add_paragraph(
        "Relevant Coursework: Data Structures & Algorithms, Machine Learning, "
        "Distributed Systems, Database Management, Operating Systems"
    )

    add_section_divider(doc)

    # ------- PROJECTS -------
    add_heading(doc, "Projects", level=2)

    p = doc.add_paragraph()
    p.add_run("1. AI Resume Parser (Personal Project)").bold = True
    doc.add_paragraph(
        "Built a Flask web app that extracts structured data from PDF/DOCX resumes "
        "using spaCy NER, PyPDF2, and Regex. Features include skill-match scoring, "
        "SQLite storage, and JSON export. Deployed on AWS EC2.\n"
        "Tech: Python, Flask, spaCy, Bootstrap, SQLite, Docker"
    )

    p = doc.add_paragraph()
    p.add_run("2. E-Commerce Recommendation Engine").bold = True
    doc.add_paragraph(
        "Designed a collaborative-filtering recommendation system using Scikit-Learn "
        "and Pandas, integrated into a Django storefront. Increased click-through rate "
        "by 28% in A/B testing.\n"
        "Tech: Python, Django, Scikit-Learn, PostgreSQL, Redis"
    )

    p = doc.add_paragraph()
    p.add_run("3. Real-Time Chat Application").bold = True
    doc.add_paragraph(
        "Full-stack chat app with Socket.IO, React, and Node.js supporting rooms, "
        "file sharing, and end-to-end encryption (AES-256).\n"
        "Tech: React, Node.js, Socket.IO, MongoDB, Docker"
    )

    add_section_divider(doc)

    # ------- CERTIFICATIONS -------
    add_heading(doc, "Certifications", level=2)
    doc.add_paragraph(
        "• AWS Certified Solutions Architect – Associate (2023)\n"
        "• Google Professional Data Engineer (2022)\n"
        "• MongoDB Certified Developer (2021)"
    )

    # Save
    output_path = "sample_resume.docx"
    doc.save(output_path)
    print(f"\n✅  Sample resume created: {os.path.abspath(output_path)}")
    print("    Upload this file at http://127.0.0.1:5000 to test the parser.\n")


if __name__ == "__main__":
    create_sample_resume()
