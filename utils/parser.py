"""
Resume Parser Utility
=====================
Handles text extraction from PDF/DOCX files and NLP-based entity extraction.
Uses spaCy for Named Entity Recognition and Regex for structured data.
"""

import re
import json

# spaCy for NLP entity extraction
# Wrapped in broad BaseException to handle Windows antivirus/OS interruptions
# during first import of newly installed packages.
nlp = None
try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
        print("[INFO] spaCy model 'en_core_web_sm' loaded successfully.")
    except OSError:
        import subprocess, sys
        print("[INFO] Downloading spaCy model 'en_core_web_sm'...")
        subprocess.run(
            [sys.executable, "-m", "spacy", "download", "en_core_web_sm"],
            check=True
        )
        nlp = spacy.load("en_core_web_sm")
        print("[INFO] spaCy model downloaded and loaded.")
except (ImportError, ModuleNotFoundError):
    print("[WARNING] spaCy not installed. Name extraction will use fallback regex.")
except BaseException as _e:
    # Catches KeyboardInterrupt and other OS-level signals that can occur
    # when Windows Defender scans newly installed packages on first import.
    print(f"[WARNING] spaCy could not be loaded ({type(_e).__name__}). "
          "Falling back to regex-based name extraction. "
          "Try restarting the server — this usually resolves on second run.")

# PDF reader
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None
    print("[WARNING] PyPDF2 not installed. PDF parsing will not work.")

# DOCX reader
try:
    from docx import Document
except ImportError:
    Document = None
    print("[WARNING] python-docx not installed. DOCX parsing will not work.")


# ==============================================================
#  COMPREHENSIVE SKILLS DATABASE
#  Add more skills here as needed
# ==============================================================
SKILLS_DB = [
    # Programming Languages
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "ruby",
    "go", "rust", "swift", "kotlin", "php", "r", "matlab", "scala", "perl",
    "bash", "shell", "powershell", "vba", "dart", "lua", "haskell", "elixir",

    # Web Frontend
    "html", "css", "react", "angular", "vue", "next.js", "nuxt.js",
    "jquery", "bootstrap", "tailwind", "sass", "less", "webpack", "vite",
    "redux", "svelte",

    # Web Backend
    "node.js", "express", "django", "flask", "fastapi", "spring", "asp.net",
    "laravel", "rails", "graphql", "rest", "soap", "api", "swagger",
    "microservices",

    # Databases
    "sql", "mysql", "postgresql", "mongodb", "sqlite", "redis", "cassandra",
    "elasticsearch", "oracle", "mssql", "dynamodb", "firebase", "neo4j",
    "mariadb", "couchdb", "influxdb",

    # Cloud & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "ci/cd",
    "terraform", "ansible", "nginx", "apache", "linux", "git", "github",
    "gitlab", "bitbucket", "circleci", "travis ci", "helm",

    # Data Science & Machine Learning
    "machine learning", "deep learning", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "matplotlib", "seaborn", "nlp",
    "computer vision", "opencv", "huggingface", "transformers", "bert",
    "gpt", "llm", "data analysis", "data visualization", "statistics",
    "hadoop", "spark", "tableau", "power bi", "data mining", "feature engineering",
    "xgboost", "lightgbm", "random forest", "neural network",

    # Mobile
    "android", "ios", "react native", "flutter", "xamarin", "ionic",

    # Testing
    "selenium", "pytest", "junit", "jest", "mocha", "cypress", "postman",
    "unit testing", "integration testing", "test automation",

    # Tools & Platforms
    "jira", "confluence", "slack", "trello", "figma", "photoshop", "excel",
    "visual studio", "vs code", "intellij", "eclipse", "jupyter", "colab",
    "linux", "unix", "windows server", "bash", "vim",

    # Methodologies
    "agile", "scrum", "kanban", "waterfall", "tdd", "bdd", "devops",
    "project management",

    # Soft Skills
    "communication", "teamwork", "leadership", "problem solving",
    "critical thinking", "time management", "adaptability",
]


# ==============================================================
#  TEXT EXTRACTION FUNCTIONS
# ==============================================================

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract raw text from a PDF file using PyPDF2.
    Iterates over all pages and concatenates the text.
    """
    if PdfReader is None:
        raise ImportError("PyPDF2 is not installed. Run: pip install PyPDF2")

    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Error reading PDF file: {str(e)}")

    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract raw text from a DOCX (Word) file using python-docx.
    Reads paragraphs and table cells.
    """
    if Document is None:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    text = ""
    try:
        doc = Document(file_path)

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # Extract text from tables (some resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text += row_text + "\n"

    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}")

    return text.strip()


# ==============================================================
#  TEXT CLEANING
# ==============================================================

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted resume text:
    - Remove extra whitespace
    - Normalize line breaks
    - Remove non-printable characters
    """
    # Remove non-printable / control characters (except newlines/tabs)
    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]', ' ', text)

    # Normalize Windows line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Collapse multiple blank lines into one
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove trailing/leading whitespace per line
    lines = [line.strip() for line in text.split('\n')]

    # Remove completely empty lines that repeat
    cleaned = []
    prev_blank = False
    for line in lines:
        if line == '':
            if not prev_blank:
                cleaned.append(line)
            prev_blank = True
        else:
            cleaned.append(line)
            prev_blank = False

    return '\n'.join(cleaned).strip()


# ==============================================================
#  ENTITY EXTRACTION FUNCTIONS
# ==============================================================

def extract_name(text: str) -> str:
    """
    Extract the full name from the resume using multiple strategies
    (ordered from most to least reliable).

    Strategies:
      1. Explicit "Name:" label in the text
      2. spaCy PERSON entity from the top of the resume
      3. First short line that looks like a human name (1-5 words, alpha only)
      4. First line of the resume (cleaned of symbols/icons)
      5. Derive a probable name from the email address
    """
    lines = text.strip().split('\n')

    # -- Common section headers / keywords to SKIP --
    skip_words = [
        'resume', 'curriculum', 'vitae', 'cv', 'profile', 'summary',
        'objective', 'address', 'phone', 'email', 'contact', 'mobile',
        'linkedin', 'github', 'portfolio', 'website', 'http', 'www',
        'education', 'experience', 'skills', 'projects', 'certifications',
        'declaration', 'references', 'date', 'place', '@',
    ]

    def _is_name_like(candidate: str) -> bool:
        """Return True if `candidate` looks like a human name."""
        candidate = candidate.strip()
        if not candidate or len(candidate) < 2:
            return False
        words = candidate.split()
        if not (1 <= len(words) <= 5):
            return False
        # Every word must start with a letter
        if not all(re.match(r'^[A-Za-z]', w) for w in words):
            return False
        # Must be mostly alphabetic (allow dots and hyphens for initials)
        alpha_chars = sum(c.isalpha() or c in '. -' for c in candidate)
        if alpha_chars / max(len(candidate), 1) < 0.85:
            return False
        # Must not be a known header/keyword
        if any(kw in candidate.lower() for kw in skip_words):
            return False
        return True

    def _clean_line(line: str) -> str:
        """Strip emojis, icons, special chars from a line, keep letters/spaces."""
        # Remove everything that's not a letter, space, dot, or hyphen
        cleaned = re.sub(r'[^A-Za-z\s\.\-]', ' ', line)
        return re.sub(r'\s{2,}', ' ', cleaned).strip()

    # -----------------------------------------------------------
    # Strategy 1: Look for explicit "Name:" label anywhere in first 15 lines
    # e.g. "Name: Anurag Shakalya" or "Full Name : John Doe"
    # -----------------------------------------------------------
    for line in lines[:15]:
        m = re.match(
            r'(?:full\s*)?name\s*[:;\-|]\s*(.+)',
            line.strip(), re.IGNORECASE
        )
        if m:
            candidate = _clean_line(m.group(1))
            if _is_name_like(candidate):
                return candidate.title()

    # -----------------------------------------------------------
    # Strategy 2: spaCy NER — scan first 10 lines for PERSON entities
    # (accepts single-word names too)
    # -----------------------------------------------------------
    first_block = '\n'.join(lines[:10])
    if nlp:
        doc = nlp(first_block)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                candidate = ent.text.strip()
                # Quick sanity: at least 2 chars, at most 5 words
                words = candidate.split()
                if 1 <= len(words) <= 5 and len(candidate) >= 2:
                    # Make sure it's not a stray keyword
                    if not any(kw in candidate.lower() for kw in skip_words):
                        return candidate

    # -----------------------------------------------------------
    # Strategy 3: First line in top-10 that is ONLY a name
    # (short line, 1-5 alphabetic words, no numbers/emails/URLs)
    # -----------------------------------------------------------
    for line in lines[:10]:
        stripped = line.strip()
        if not stripped:
            continue
        cleaned = _clean_line(stripped)
        if _is_name_like(cleaned) and len(cleaned) <= 50:
            return cleaned.title()

    # -----------------------------------------------------------
    # Strategy 4: Take the very first non-empty line, strip all
    # non-alpha chars, and check if the result looks like a name
    # -----------------------------------------------------------
    for line in lines[:5]:
        stripped = line.strip()
        if not stripped:
            continue
        cleaned = _clean_line(stripped)
        # Remove very short fragments (single chars from icons)
        words = [w for w in cleaned.split() if len(w) >= 2]
        candidate = ' '.join(words)
        if candidate and 1 <= len(words) <= 5:
            if not any(kw in candidate.lower() for kw in skip_words):
                return candidate.title()
        break  # only try the first non-empty line

    # -----------------------------------------------------------
    # Strategy 5: Derive name from the email address
    # e.g. "anuragshakalya@gmail.com" -> "Anuragshakalya"
    # e.g. "john.doe@company.com"     -> "John Doe"
    # -----------------------------------------------------------
    email = extract_email(text)
    if email and email != "Not Found":
        local_part = email.split('@')[0]
        # Remove numbers and common prefixes
        local_part = re.sub(r'[0-9_]', ' ', local_part)
        # Split on dots and hyphens
        parts = re.split(r'[\.\-\+]+', local_part)
        # Filter out very short fragments
        parts = [p.strip() for p in parts if len(p.strip()) >= 2]
        if parts:
            return ' '.join(p.capitalize() for p in parts)

    return "Not Found"


def extract_email(text: str) -> str:
    """
    Extract email address using a robust RFC-compliant-ish regex.
    Returns the first match found.
    """
    pattern = r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,7}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else "Not Found"


def extract_phone(text: str) -> str:
    """
    Extract phone number using multiple regex patterns to handle
    various international and local formats.
    """
    patterns = [
        # +91 98765 43210  or  +1 (555) 123-4567
        r'\+\d{1,3}[\s\-\.]?\(?\d{1,4}\)?[\s\-\.]?\d{3,5}[\s\-\.]?\d{4,6}',
        # (555) 123-4567
        r'\(\d{3}\)[\s\-\.]?\d{3}[\s\-\.]?\d{4}',
        # 555-123-4567  or  555.123.4567  or  5551234567
        r'\b\d{3}[\s\-\.]\d{3}[\s\-\.]\d{4}\b',
        # 10-digit number (India mobile)
        r'\b[6-9]\d{9}\b',
        # General 10-15 digit number
        r'\b\d{10,15}\b',
    ]

    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Take the first match and clean extra whitespace
            phone = re.sub(r'\s+', ' ', matches[0]).strip()
            # Sanity check: must have at least 10 digits
            digits = re.sub(r'\D', '', phone)
            if len(digits) >= 10:
                return phone

    return "Not Found"


def extract_linkedin(text: str) -> str:
    """
    Extract the FULL LinkedIn profile URL.
    Handles variations like:
      - https://www.linkedin.com/in/username
      - linkedin.com/in/username-123abc
      - www.linkedin.com/in/username_name
    Ensures the URL is not truncated.
    """
    # Match the full LinkedIn profile URL including query params and trailing paths
    pattern = (
        r'(?:https?://)?'           # optional scheme
        r'(?:www\.)?'               # optional www
        r'linkedin\.com/in/'        # base path
        r'[\w\-\.%]+'              # username (allows hyphens, dots, percent-encoding)
        r'(?:/[\w\-\.%]*)*'        # optional sub-paths
        r'(?:\?[^\s<>"]*)?'        # optional query string
    )

    matches = re.findall(pattern, text, re.IGNORECASE)

    if matches:
        url = matches[0].strip().rstrip('.,;)')  # remove trailing punctuation

        # Ensure the URL has https:// prefix
        if not url.lower().startswith('http'):
            url = 'https://' + url

        return url

    return "Not Found"


def extract_github(text: str) -> str:
    """
    Extract the GitHub profile URL.
    Handles formats like:
      - https://github.com/username
      - github.com/username
      - www.github.com/username
    Also supports fallback from labels like:
      - GitHub: username
      - Github - username
    """
    # Full URL pattern (including optional trailing paths/query)
    url_pattern = (
        r'(?:https?://)?'
        r'(?:www\.)?'
        r'github\.com/'
        r'[A-Za-z0-9](?:[A-Za-z0-9\-]{0,38})'
        r'(?:/[\w\-\.%]*)*'
        r'(?:\?[^\s<>\"]*)?'
    )

    matches = re.findall(url_pattern, text, re.IGNORECASE)
    if matches:
        url = matches[0].strip().rstrip('.,;)')
        if not url.lower().startswith('http'):
            url = 'https://' + url
        return url

    # Fallback: GitHub label + username
    label_pattern = r'github\s*[:\-|]\s*([A-Za-z0-9][A-Za-z0-9\-]{0,38})'
    label_match = re.search(label_pattern, text, re.IGNORECASE)
    if label_match:
        username = label_match.group(1)
        return f"https://github.com/{username}"

    return "Not Found"


def extract_skills(text: str) -> list:
    """
    Extract skills from resume text by matching against the SKILLS_DB list.
    - Case-insensitive matching
    - Uses word boundaries for short skill names to avoid false positives
    - Removes duplicate skills
    Returns a sorted, deduplicated list of found skills.
    """
    text_lower = text.lower()
    found_skills = []

    for skill in SKILLS_DB:
        skill_lower = skill.lower()

        # Use word boundaries for short skills (e.g., 'c', 'r', 'go')
        # to avoid matching them inside longer words
        if len(skill) <= 3:
            pattern = r'(?<![a-zA-Z])' + re.escape(skill_lower) + r'(?![a-zA-Z])'
        else:
            # For multi-word or longer skills, simple substring search is fine
            pattern = re.escape(skill_lower)

        if re.search(pattern, text_lower):
            # Format the skill nicely for display
            if '.' in skill or '/' in skill or skill.upper() in ['AWS', 'GCP', 'SQL', 'CSS',
                                                                   'HTML', 'API', 'NLP', 'TDD',
                                                                   'BDD', 'VBA']:
                found_skills.append(skill)  # keep original casing for acronyms/special
            elif len(skill) <= 4 and skill.upper() == skill.upper():
                found_skills.append(skill.upper())
            else:
                found_skills.append(skill.title())

    # Deduplicate while preserving order (case-insensitive dedup)
    seen = set()
    unique_skills = []
    for skill in found_skills:
        key = skill.lower()
        if key not in seen:
            seen.add(key)
            unique_skills.append(skill)

    return sorted(unique_skills)


def _extract_section(text: str, section_keywords: list) -> str:
    """
    Internal helper: extract the content of a named section from resume text.

    How it works:
    - Scans line-by-line looking for a section header matching `section_keywords`
    - Collects lines until it hits a line that looks like a DIFFERENT section header
    - Returns up to 40 lines of section content
    """
    # All known section header keywords (used to detect end of current section)
    all_section_headers = [
        'experience', 'education', 'skills', 'projects', 'certifications',
        'awards', 'publications', 'references', 'summary', 'objective',
        'work history', 'employment', 'academic', 'achievements', 'interests',
        'languages', 'volunteer', 'activities', 'contact', 'personal',
        'internship', 'training', 'courses', 'hobbies', 'declaration',
    ]

    # The "stop" headers are all headers EXCEPT the ones we're looking for
    section_kws_lower = [kw.lower() for kw in section_keywords]
    stop_keywords = [h for h in all_section_headers if h not in section_kws_lower]

    lines = text.split('\n')
    section_content = []
    in_section = False
    empty_line_count = 0

    for line in lines:
        stripped = line.strip()
        line_lower = stripped.lower()

        # ---- Detect section start ----
        if not in_section:
            if any(kw.lower() in line_lower for kw in section_keywords):
                # Validate it's actually a header (short line, not a sentence)
                if len(stripped) < 80:
                    in_section = True
            continue  # skip the header line itself

        # ---- Detect section end ----
        if in_section:
            # Stop if we encounter another section header
            if any(kw in line_lower for kw in stop_keywords) and len(stripped) < 60:
                break

            # Stop after too many consecutive blank lines (new section likely started)
            if stripped == '':
                empty_line_count += 1
                if empty_line_count >= 3:
                    break
            else:
                empty_line_count = 0
                section_content.append(stripped)

            # Hard cap at 40 lines to avoid runaway extraction
            if len(section_content) >= 40:
                break

    return '\n'.join(section_content) if section_content else "Not Found"


def extract_education(text: str) -> str:
    """
    Extract the Education section.
    Falls back to regex-based degree detection if no section header is found.
    """
    keywords = ['education', 'academic background', 'qualifications',
                 'academics', 'educational qualification']
    content = _extract_section(text, keywords)

    if content == "Not Found":
        # Fallback: look for degree keywords anywhere in the text
        degree_pattern = (
            r'(?:B\.?Tech|B\.?E\.?|B\.?Sc\.?|B\.?Com\.?|B\.?A\.?|'
            r'M\.?Tech|M\.?E\.?|M\.?Sc\.?|M\.?B\.?A\.?|M\.?C\.?A\.?|'
            r'Ph\.?D\.?|Bachelor|Master|Doctor|Diploma|High School|'
            r'10th|12th|Secondary|Senior Secondary)'
            r'[^\n]{0,120}'
        )
        matches = re.findall(degree_pattern, text, re.IGNORECASE)
        if matches:
            return '\n'.join(dict.fromkeys(m.strip() for m in matches[:6]))

    return content


def extract_experience(text: str) -> str:
    """
    Extract the Work Experience section.
    Looks for multiple common header variants.
    """
    keywords = [
        'work experience', 'professional experience', 'experience',
        'employment history', 'work history', 'employment',
        'career history', 'internship', 'internships',
    ]
    return _extract_section(text, keywords)


def extract_projects(text: str) -> str:
    """
    Extract the Projects section.
    """
    keywords = [
        'projects', 'personal projects', 'academic projects',
        'project work', 'key projects', 'notable projects',
    ]
    return _extract_section(text, keywords)


# ==============================================================
#  SKILL MATCH ANALYSIS
# ==============================================================

def calculate_skill_match(resume_skills: list, job_description: str) -> tuple:
    """
    Calculate how well the candidate's skills match a given job description.

    Returns:
        percentage (float)  : match score 0-100
        matched (list)      : skills present in both resume and job description
        missing (list)      : skills in JD but not in resume
    """
    if not job_description or not resume_skills:
        return 0.0, [], []

    # Extract skills from the job description using the same extractor
    job_skills = extract_skills(job_description)

    if not job_skills:
        return 0.0, [], []

    resume_lower = {s.lower() for s in resume_skills}
    job_lower = {s.lower() for s in job_skills}

    matched_lower = resume_lower & job_lower
    missing_lower = job_lower - resume_lower

    # Map back to display-friendly casing
    matched = [s for s in job_skills if s.lower() in matched_lower]
    missing = [s for s in job_skills if s.lower() in missing_lower]

    percentage = round((len(matched) / len(job_skills)) * 100, 1)

    return percentage, matched, missing


# ==============================================================
#  MAIN ENTRY POINT
# ==============================================================

def parse_resume(file_path: str, file_extension: str, job_description: str = "") -> dict:
    """
    Main function: extract all resume information from a PDF or DOCX file.

    Args:
        file_path       : Absolute path to the uploaded resume file
        file_extension  : '.pdf' or '.docx'/'.doc'
        job_description : Optional job description for skill-match scoring

    Returns:
        dict with all extracted fields
    """
    # Step 1 — Extract raw text based on file type
    ext = file_extension.lower()
    if ext == '.pdf':
        raw_text = extract_text_from_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file format '{ext}'. Please upload a PDF or DOCX file."
        )

    if not raw_text.strip():
        raise ValueError(
            "Could not extract any text from the file. "
            "It may be an image-based PDF or a corrupted file."
        )

    # Step 2 — Clean & normalize text
    cleaned_text = clean_text(raw_text)

    # Step 3 — Extract individual fields
    skills = extract_skills(cleaned_text)
    match_pct, matched_skills, missing_skills = calculate_skill_match(
        skills, job_description
    )

    parsed = {
        "name":       extract_name(cleaned_text),
        "email":      extract_email(cleaned_text),
        "phone":      extract_phone(cleaned_text),
        "linkedin":   extract_linkedin(cleaned_text),
        "github":     extract_github(cleaned_text),
        "skills":     skills,
        "education":  extract_education(cleaned_text),
        "experience": extract_experience(cleaned_text),
        "projects":   extract_projects(cleaned_text),
        "skill_match": {
            "percentage":       match_pct,
            "matched_skills":   matched_skills,
            "missing_skills":   missing_skills,
            "total_jd_skills":  len(matched_skills) + len(missing_skills),
            "matched_count":    len(matched_skills),
        },
        "raw_text_preview": (
            cleaned_text[:800] + "\n...[truncated]"
            if len(cleaned_text) > 800
            else cleaned_text
        ),
    }

    return parsed
